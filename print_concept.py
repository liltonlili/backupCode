#coding:utf-8
import pymongo
from docx import Document


mongo_url = "localhost"
mongodb = pymongo.MongoClient(mongo_url)

fHandler = open(u"./全概念/concept.txt", 'wb')
docHandler = Document()
concept_lists = mongodb.stock.myconcept.distinct("concept")
for concept in concept_lists:
    results = mongodb.stock.myconcept.find({"concept":concept})
    if results.count() > 0:
        tcount = 1
        for result in results:
            concept = result['concept']
            stockid = result['stockid']
            stockname = result['name']
            reason = result['reason']
            text = "[%s %s/%s] %s  %s:\n## reason:\n%s\n" % (concept, tcount, results.count(), stockid, stockname, reason)
            docHandler.add_paragraph("[%s %s/%s] %s  %s:" %(concept, tcount, results.count(), stockid, stockname))
            docHandler.add_paragraph("## reason:")
            docHandler.add_paragraph("%s" %reason)
            fHandler.write(text.encode("utf-8"))
            tcount += 1
docHandler.save(u"./全概念/concept.docx")
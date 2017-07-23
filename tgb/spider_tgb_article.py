# coding:utf-8
import requests
import time as tt
from docx import Document
from docx.shared import Inches
from lxml import etree
import common
import os
import sys
import logging
import logging.config
import spider_tgb
reload(sys)
sys.setdefaultencoding("utf8")

'''
爬取某V在某贴下的所有发言
'''


def get_total_page(root):
    tpage = 0
    obj_tpage = root.xpath('//div[@class="left t_page01"]/span[3]/text()')
    if len(obj_tpage)>0:
        tpage = obj_tpage[1].replace(u"/", u"").replace(u'页', "")
        logging.getLogger().info("get total page of:%s" % tpage)
    else:
        logging.getLogger().error("can't get the tpage, len = 0")
    return tpage


# 爬取某一页
# 返回total_page_no, status
def crawl_page(targetUrl, current_page):
    global s, pic_count, docHandler, mydir, txtHandler
    crawl_status = False
    total_page = 0
    try:
        retry_count = 20
        logging.getLogger().info("Crawl url:%s"%targetUrl)

        # 发送请求
        while retry_count > 0:
            try:
                r_tmp = s.get(targetUrl, headers=spider_tgb.general_headers)
                print "get target url:%s" % targetUrl
                if r_tmp.status_code == 200:
                    status = True
                    logging.getLogger().info("Succeed to connect url with repeat:%s" % (retry_count-20))
                    break
            except:
                pass
            tt.sleep(10)
            retry_count -= 1

        # 创建xpath对象
        content = r_tmp.content.decode('utf8').replace(u'\x00', u'')
        root = etree.HTML(content)

        # 解析总页码
        total_page = get_total_page(root)   # 总页码

        # 解析具体内容
        if current_page == 1:   #页码为1，解析发帖主内容
            parse_main_content(root)

        parse_update_content(root)  # 解析随后的跟帖内容
        crawl_status = True
        logging.getLogger().info("Succeed in crawl_page, url:%s" %targetUrl)
    except Exception, err:
        logging.getLogger().error("Error when crawl_page, url:%s, error:%s" % (targetUrl, err))
    return crawl_status, total_page

# 解析跟帖内容
def parse_update_content(root):
    global pic_count, docHandler, txtHandler, user_id
    objs = root.xpath('//div[@class="pc_p_nr user_%s"]' %user_id)
    for obj in objs:
        time_str = obj.xpath('./div[@class="pc_yc_an"]/div[@class="left pcyc_l"]/span/text()')[0]
        time_date = time_str.split(" ")[0]

        # 文字内容
        main_content = obj.xpath('./div[@class="pcnr_wz"]/p')[0].xpath("string(.)")

        prefix = u"\n----%s----:" % (time_str)
        docHandler.add_paragraph("%s\n\n"%prefix)


        mainContent = main_content.replace(" ","").replace(u'\xa0','\n').replace("\t","").replace("\n\n","\n").replace("\n\n","\n").replace("\n\n","\n").replace("\r\n","")
        mainContent = mainContent.replace(u"[淘股吧]", u"")
        mainContent = mainContent.split('原帖由')[0]
        #引用
        try:
            referContent = obj.xpath('./div[@class="pcnr_wz"]/p/span')[1].xpath('string(.)').replace(u" ", u"").replace(u' ', u'')
            referContent = referContent.replace(u"发表",u"发表\n\t")
            referContent = referContent.replace(u'\xa0','').replace(" ","").strip()
        except:
            referContent = 'None'

        mainContents = "\t"+mainContent
        textContent = "%s\n\n##%s" %(mainContents, referContent)
        try:
            docHandler.add_paragraph(textContent)
        except:
            try:
                docHandler.add_paragraph(textContent.decode("utf8"))
            except:
                docHandler.add_paragraph(textContent.decode("gbk"))

        # 图片
        pics = obj.xpath('./div[@class="pcnr_wz"]/div[@align="center"]/img/@data-original')

        for pic in pics:
            if len(pic) <= 5:
                continue
            if '.jpg' in pic:
                file_name = u"%s_%s.jpg" % (time_date, pic_count)
            elif ".png" in pic:
                file_name = u"%s_%s.png" % (time_date, pic_count)
            elif ".bmp" in pic:
                file_name = u"%s_%s.bmp" % (time_date, pic_count)
            else:
                file_name = u"%s_%s.jpg" % (time_date, pic_count)

            # 图片名
            file_name = os.path.join(mydir, u"pic/%s"%file_name)
            try:
                with open(file_name, 'wb') as fhandler:
                    rPic = s.get(url = pic, headers=spider_tgb.image_headers)
                    print "get pic url:%s" % pic
                    fhandler.write(rPic.content)
                docHandler.add_picture(file_name,width=Inches(6))
            except Exception,err:
                    logging.getLogger().error("Error when add picture, err:%s, file:%s" %(err, file_name))
            finally:
                pic_count += 1
# 解析主帖内容
def parse_main_content(root):
    global pic_count, docHandler, txtHandler
    time_str = root.xpath('//div[@class="p_wenz"]//span[@class="p_tatime"]/text()')[0]
    time_date = time_str.split(" ")[0]

    # 正文文字
    main_content = root.xpath('//div[@class="p_wenz"]//div[@class="p_coten"]')[0].xpath("string(.)")

    prefix = u"\n----%s----:" % (time_str)
    docHandler.add_paragraph("%s\n\n"%prefix)

    mainContent = main_content.replace(u" ",u"").replace(u'\xa0',u'\n').replace(u"\t",u"").replace(u"\n\n",u"\n").replace(u"\n\n",u"\n").replace(u"\n\n",u"\n").replace(u"\r\n",u"")
    mainContents = "\t"+mainContent
    try:
        docHandler.add_paragraph(mainContents)
    except:
        try:
            docHandler.add_paragraph(mainContents.decode("utf8"))
        except:
            docHandler.add_paragraph(mainContents.decode("gbk"))

    # 图片
    pics = root.xpath('//div[@class="p_wenz"]//div[@class="p_coten"]/div[@align="center"]/img/@data-original')

    for pic in pics:
        if len(pic) <= 5:
            continue
        if '.jpg' in pic:
            file_name = u"%s_%s.jpg" % (time_date, pic_count)
        elif ".png" in pic:
            file_name = u"%s_%s.png" % (time_date, pic_count)
        elif ".bmp" in pic:
            file_name = u"%s_%s.bmp" % (time_date, pic_count)
        else:
            file_name = u"%s_%s.jpg" % (time_date, pic_count)

        # 图片名
        file_name = os.path.join(mydir, u"pic/%s"%file_name)
        try:
            with open(file_name, 'wb') as fhandler:
                rPic = s.get(url = pic, headers=spider_tgb.image_headers)
                print "get pic url: %s" % pic
                fhandler.write(rPic.content)
            docHandler.add_picture(file_name,width=Inches(6))
        except Exception,err:
                logging.getLogger().error("Error when add picture, err:%s, file:%s" %(err, file_name))
        finally:
            pic_count += 1

if __name__ == "__main__":
    global s, baseUrl, docHandler, txtHandler, mydir, pic_count

    article_id = '1449165' # 文章的ID
    user_id = '591447'    #User的ID
    id_Name = u'刺客实盘_著名刺客'     # 文章名字 + 跟踪者名字
    log_id_name = u'cksp_zmck'

    logdir = os.path.join(u"D:/Money/lilton_code/Market_Mode/learnModule/logs/articles/%s"%log_id_name,"")
    mydir = u'D:/Money/lilton_code/Market_Mode/learnModule/article/%s'%id_Name
    pic_count = 0
    common.look_up_dir(logdir)
    common.look_up_dir(os.path.join(mydir, "pic"))
    logging.config.fileConfig("./conf/conf_log.txt", defaults={'logdir': logdir})
    logging.getLogger().info("report main startup")
    baseUrl = "http://www.taoguba.com.cn/lookUserTopic?pageNo=PAGENO&topicID=ARTID&lookUserID=USERID"

    txtHandler = open(os.path.join(mydir,"record.txt"),'w')     # 用来记录内容的txt
    docHandler = Document()             # doc记录

    # 开始登陆淘股吧
    s = requests.Session()
    login_url = 'http://www.taoguba.com.cn/newLogin'
    r = s.post(url=login_url, headers=spider_tgb.headers, data=spider_tgb.data, allow_redirects=True)

    page_no = 10    # 默认有10页
    for i in range(1, page_no):
        targetUrl = baseUrl.replace("PAGENO", str(i)).replace("ARTID", article_id).replace("USERID", user_id) # 实际访问的页面
        status, total_page_no= crawl_page(targetUrl, i)
        if i >= int(total_page_no):
            break
        i += 1

    docHandler.save(os.path.join(mydir, "%s.docx" % id_Name))